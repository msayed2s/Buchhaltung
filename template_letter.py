"""
Vorlagen-System für Anschreiben (Rechnungen & Fahrtkosten).

Füllt die Platzhalter der Datei ``Vorlage.docx`` im Projekt-/Programmverzeichnis:

    [Name des Empfängers]   -> Empfängername (aus Kundendaten)
    [Straße und Hausnummer] -> Adresse Zeile 1 (aus Kundendaten)
    [PLZ Ort]               -> Adresse Zeile 2 (aus Kundendaten)
    [Datum]                 -> Erstellungsdatum (deutsches Format)
    [Betreff]               -> Betreff (Standard aus Vorlage.properties, überschreibbar)
    [Text des Schreibens]   -> Fließtext (Standard aus Vorlage.properties, überschreibbar)

Die Standardwerte für Betreff und Text liegen in ``Vorlage.properties`` und
können dort jederzeit ausgetauscht werden.
"""
import os
import sys
import subprocess

from docx import Document
from docx.oxml import OxmlElement


# ---------- Pfade (funktioniert im Dev-Modus und als PyInstaller-.exe) ----------
def base_dir():
    if getattr(sys, "frozen", False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))


def _find_file(name):
    """Sucht eine Datei neben der .exe/dem Skript und im PyInstaller-Bundle."""
    candidates = [os.path.join(base_dir(), name)]
    meipass = getattr(sys, "_MEIPASS", None)
    if meipass:
        candidates.append(os.path.join(meipass, name))
    for c in candidates:
        if os.path.exists(c):
            return c
    return candidates[0]


def vorlage_path():
    return _find_file("Vorlage.docx")


def properties_path():
    return _find_file("Vorlage.properties")


def ensure_external_files():
    """Kopiert die gebündelten Vorlage-Dateien neben die .exe, falls sie dort
    noch nicht liegen — so kann der Nutzer Vorlage.docx/.properties bearbeiten."""
    meipass = getattr(sys, "_MEIPASS", None)
    if not meipass:
        return
    import shutil
    for name in ("Vorlage.docx", "Vorlage.properties"):
        target = os.path.join(base_dir(), name)
        source = os.path.join(meipass, name)
        if not os.path.exists(target) and os.path.exists(source):
            try:
                shutil.copy2(source, target)
            except OSError:
                pass


# ---------- .properties ----------
_DEFAULTS = {
    "betreff": "Rechnung für erbrachte Leistungen",
    "text": "vielen Dank für Ihren Auftrag.\nFür die erbrachten Leistungen erlaube ich mir, wie folgt zu berechnen:",
    "betreff_fahrten": "Fahrtkostenaufstellung",
    "text_fahrten": "nachfolgend erhalten Sie die Aufstellung der betrieblich veranlassten Fahrten:",
}


def load_properties():
    """Liest Vorlage.properties (key=value, # als Kommentar, \\n als Zeilenumbruch)."""
    props = dict(_DEFAULTS)
    path = properties_path()
    if not os.path.exists(path):
        return props
    try:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.rstrip("\n")
                if not line.strip() or line.lstrip().startswith("#"):
                    continue
                if "=" not in line:
                    continue
                key, val = line.split("=", 1)
                props[key.strip()] = val.replace("\\n", "\n")
    except OSError:
        pass
    return props


# ---------- docx-Platzhalter ersetzen ----------
def _replace_in_paragraph(paragraph, mapping):
    full = "".join(run.text for run in paragraph.runs)
    if not any(k in full for k in mapping):
        return
    for k, v in mapping.items():
        full = full.replace(k, v)
    if paragraph.runs:
        paragraph.runs[0].text = full
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(full)


def _insert_paragraph_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, paragraph._parent)
    try:
        new_para.style = paragraph.style
    except Exception:
        pass
    if text:
        new_para.add_run(text)
    return new_para


def _all_paragraphs(doc):
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
    for section in doc.sections:
        for hf in (section.header, section.footer):
            yield from hf.paragraphs


def fill_template(values, out_docx):
    """Füllt Vorlage.docx mit ``values`` und speichert nach ``out_docx``.

    ``values`` erwartet die Schlüssel: name, strasse, plz_ort, datum,
    betreff, text (Zeilen mit \\n), optional anrede.
    """
    doc = Document(vorlage_path())
    body_text = values.get("text", "") or ""
    single = {
        "[Name des Empfängers]": values.get("name", ""),
        "[Straße und Hausnummer]": values.get("strasse", ""),
        "[PLZ Ort]": values.get("plz_ort", ""),
        "[Datum]": values.get("datum", ""),
        "[Betreff]": values.get("betreff", ""),
    }
    anrede = values.get("anrede")
    if anrede:
        single["Sehr geehrte Damen und Herren,"] = anrede + ","

    for para in list(_all_paragraphs(doc)):
        if "[Text des Schreibens]" in para.text:
            lines = body_text.split("\n")
            _replace_in_paragraph(para, {"[Text des Schreibens]": lines[0] if lines else ""})
            ref = para
            for line in lines[1:]:
                ref = _insert_paragraph_after(ref, line)
        else:
            _replace_in_paragraph(para, single)

    doc.save(out_docx)
    return out_docx


# ---------- docx -> pdf (best effort, offline) ----------
def convert_to_pdf(docx_path, pdf_path):
    """Wandelt eine .docx in eine .pdf um.

    Nutzt bevorzugt Microsoft Word (direkt über COM), sonst LibreOffice.
    Wirft RuntimeError, wenn keine Konvertierung möglich ist.
    """
    docx_path = os.path.abspath(docx_path)
    pdf_path = os.path.abspath(pdf_path)
    errors = []

    # 1) Microsoft Word direkt über COM (robuster als docx2pdf)
    try:
        return _word_to_pdf(docx_path, pdf_path)
    except Exception as e:
        errors.append(f"Word: {e}")

    # 2) LibreOffice
    soffice = _find_soffice()
    if soffice:
        out_dir = os.path.dirname(pdf_path)
        try:
            subprocess.run([soffice, "--headless", "--convert-to", "pdf",
                            "--outdir", out_dir, docx_path],
                           check=True, capture_output=True, timeout=120)
            produced = os.path.join(out_dir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
            if os.path.exists(produced):
                if os.path.abspath(produced) != pdf_path:
                    os.replace(produced, pdf_path)
                return pdf_path
        except Exception as e:
            errors.append(f"LibreOffice: {e}")

    raise RuntimeError(
        "Die PDF-Umwandlung benötigt Microsoft Word oder LibreOffice.\n"
        "Alternativ steht die Word-Datei (.docx) zur Verfügung.\n\n"
        + "\n".join(errors))


def _word_to_pdf(docx_path, pdf_path):
    """Konvertiert per Microsoft Word COM. Wird von convert_to_pdf genutzt."""
    import pythoncom
    import win32com.client

    wdFormatPDF = 17
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(docx_path, ReadOnly=1)
        try:
            # bevorzugt: ExportAsFixedFormat (zuverlässig)
            doc.ExportAsFixedFormat(pdf_path, ExportFormat=wdFormatPDF)
        except Exception:
            doc.SaveAs(pdf_path, FileFormat=wdFormatPDF)
        return pdf_path
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


def _find_soffice():
    for cand in ("soffice", "soffice.exe",
                 r"C:\Program Files\LibreOffice\program\soffice.exe",
                 r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"):
        if os.path.isabs(cand) and os.path.exists(cand):
            return cand
        from shutil import which
        found = which(cand)
        if found:
            return found
    return None
